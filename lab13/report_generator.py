from docx import Document
from docx.shared import Inches
import openpyxl

class ReportGenerator:
    def generate_doc_report(self, filename, item, size, fabric_consumption, sewing_cost):
        """Создает отчет в формате .doc"""
        document = Document()
        document.add_heading('Отчет о расходе ткани и стоимости пошива', 0)
        document.add_paragraph(f'Изделие: {item}')
        document.add_paragraph(f'Размер: {size}')
        document.add_paragraph(f'Расход ткани: {fabric_consumption} м')
        document.add_paragraph(f'Стоимость пошива: {sewing_cost} руб.')
        document.save(filename)

    def generate_xls_report(self, filename, item, size, fabric_consumption, sewing_cost):
        """Создает отчет в формате .xls"""
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet['A1'] = 'Изделие'
        sheet['B1'] = 'Размер'
        sheet['C1'] = 'Расход ткани'
        sheet['D1'] = 'Стоимость пошива'
        sheet['A2'] = item
        sheet['B2'] = size
        sheet['C2'] = fabric_consumption
        sheet['D2'] = sewing_cost
        workbook.save(filename)